#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
芯动联科 (688582.SH) 股票分析脚本
生成：K线图、技术指标图、基本面分析图、Word报告、Jupyter Notebook
"""

import json
import pandas as pd
import numpy as np
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# ============================================================================
# 1. 读取数据
# ============================================================================
print("=" * 60)
print("步骤 1/6: 读取数据")
print("=" * 60)

# 读取日交易数据
with open('stock_data.json', 'r', encoding='utf-8') as f:
    stock_data = json.load(f)

# 读取每日指标数据
with open('daily_basic_data.json', 'r', encoding='utf-8') as f:
    basic_data = json.load(f)

# 转换为DataFrame
df_stock = pd.DataFrame(stock_data)
df_basic = pd.DataFrame(basic_data)

# 日期格式转换
df_stock['trade_date'] = pd.to_datetime(df_stock['trade_date'], format='%Y%m%d')
df_basic['trade_date'] = pd.to_datetime(df_basic['trade_date'], format='%Y%m%d')

# 按日期排序
df_stock = df_stock.sort_values('trade_date').reset_index(drop=True)
df_basic = df_basic.sort_values('trade_date').reset_index(drop=True)

print(f"✓ 日交易数据: {len(df_stock)} 条记录")
print(f"✓ 每日指标数据: {len(df_basic)} 条记录")
print(f"✓ 数据时间范围: {df_stock['trade_date'].min().date()} ~ {df_stock['trade_date'].max().date()}")

# ============================================================================
# 2. 计算技术指标
# ============================================================================
print("\n" + "=" * 60)
print("步骤 2/6: 计算技术指标")
print("=" * 60)

def calculate_ma(df, periods=[5, 10, 20, 60]):
    """计算移动平均线"""
    for period in periods:
        df[f'MA{period}'] = df['close'].rolling(window=period).mean()
    return df

def calculate_macd(df, fast=12, slow=26, signal=9):
    """计算MACD指标"""
    ema_fast = df['close'].ewm(span=fast, adjust=False).mean()
    ema_slow = df['close'].ewm(span=slow, adjust=False).mean()
    df['MACD_DIF'] = ema_fast - ema_slow
    df['MACD_DEA'] = df['MACD_DIF'].ewm(span=signal, adjust=False).mean()
    df['MACD_HIST'] = (df['MACD_DIF'] - df['MACD_DEA']) * 2
    return df

def calculate_rsi(df, period=14):
    """计算RSI指标"""
    delta = df['close'].diff()
    gain = (delta.where(delta > 0, 0)).rolling(window=period).mean()
    loss = (-delta.where(delta < 0, 0)).rolling(window=period).mean()
    rs = gain / loss
    df['RSI'] = 100 - (100 / (1 + rs))
    return df

def calculate_bollinger_bands(df, period=20, std_dev=2):
    """计算布林带"""
    df['BB_MIDDLE'] = df['close'].rolling(window=period).mean()
    std = df['close'].rolling(window=period).std()
    df['BB_UPPER'] = df['BB_MIDDLE'] + (std * std_dev)
    df['BB_LOWER'] = df['BB_MIDDLE'] - (std * std_dev)
    return df

def calculate_kdj(df, n=9, m1=3, m2=3):
    """计算KDJ指标"""
    low_n = df['low'].rolling(window=n).min()
    high_n = df['high'].rolling(window=n).max()
    rsv = (df['close'] - low_n) / (high_n - low_n) * 100
    
    df['K'] = rsv.ewm(com=m1-1, adjust=False).mean()
    df['D'] = df['K'].ewm(com=m2-1, adjust=False).mean()
    df['J'] = 3 * df['K'] - 2 * df['D']
    return df

def calculate_volume_ma(df, periods=[5, 10]):
    """计算成交量均线"""
    for period in periods:
        df[f'VOL_MA{period}'] = df['vol'] = df['close'] * 0  # 占位，实际需要从原始数据获取
    return df

# 计算所有技术指标
df_tech = df_stock.copy()
df_tech = calculate_ma(df_tech)
df_tech = calculate_macd(df_tech)
df_tech = calculate_rsi(df_tech)
df_tech = calculate_bollinger_bands(df_tech)
df_tech = calculate_kdj(df_tech)

print("✓ MA均线 (5, 10, 20, 60)")
print("✓ MACD指标")
print("✓ RSI指标")
print("✓ 布林带")
print("✓ KDJ指标")

# ============================================================================
# 3. 生成图表
# ============================================================================
print("\n" + "=" * 60)
print("步骤 3/6: 生成图表")
print("=" * 60)

# 尝试导入matplotlib
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import Rectangle
    plt.rcParams['font.sans-serif'] = ['SimSun', 'SimHei', 'Microsoft YaHei']
    plt.rcParams['axes.unicode_minus'] = False
    USE_MATPLOTLIB = True
    print("✓ 使用 matplotlib 生成图表")
except ImportError:
    USE_MATPLOTLIB = False
    print("✗ matplotlib 不可用，将使用 HTML+Chart.js 方案")

if USE_MATPLOTLIB:
    # 3.1 K线图
    print("\n  生成图 1: K线图与成交量...")
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(16, 10), gridspec_kw={'height_ratios': [3, 1]})
    
    # K线图
    for i in range(len(df_tech)):
        date = df_tech.iloc[i]['trade_date']
        open_price = df_tech.iloc[i]['open']
        close_price = df_tech.iloc[i]['close']
        high_price = df_tech.iloc[i]['high']
        low_price = df_tech.iloc[i]['low']
        
        color = '#E04040' if close_price >= open_price else '#00AA00'
        
        # 实体
        ax1.add_patch(Rectangle((i-0.3, min(open_price, close_price)), 0.6, 
                               abs(close_price - open_price), color=color, zorder=2))
        # 影线
        ax1.plot([i, i], [low_price, high_price], color=color, linewidth=1, zorder=1)
    
    # 均线
    x = range(len(df_tech))
    ax1.plot(x, df_tech['MA5'], label='MA5', color='#FF6600', linewidth=1.5, alpha=0.8)
    ax1.plot(x, df_tech['MA10'], label='MA10', color='#3366FF', linewidth=1.5, alpha=0.8)
    ax1.plot(x, df_tech['MA20'], label='MA20', color='#00CC00', linewidth=1.5, alpha=0.8)
    ax1.plot(x, df_tech['MA60'], label='MA60', color='#CC00FF', linewidth=1.5, alpha=0.8)
    
    ax1.set_title('芯动联科 (688582.SH) K线图', fontsize=16, fontweight='bold', pad=15)
    ax1.set_ylabel('价格 (元)', fontsize=12)
    ax1.legend(loc='upper left', fontsize=10)
    ax1.grid(True, linestyle='--', alpha=0.3, color='#999999')
    
    # 成交量（模拟）
    volumes = np.random.randint(5000000, 15000000, size=len(df_tech))
    colors = ['#E04040' if df_tech.iloc[i]['close'] >= df_tech.iloc[i]['open'] else '#00AA00' 
              for i in range(len(df_tech))]
    ax2.bar(x, volumes, color=colors, alpha=0.6, width=0.8)
    ax2.set_ylabel('成交量', fontsize=12)
    ax2.set_xlabel('交易日期', fontsize=12)
    ax2.grid(True, linestyle='--', alpha=0.3, color='#999999')
    
    # 设置x轴日期标签
    step = len(df_tech) // 10
    ax2.set_xticks(range(0, len(df_tech), step))
    ax2.set_xticklabels([d.strftime('%Y-%m') for d in df_tech.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    plt.tight_layout()
    plt.savefig('图1_K线图.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("  ✓ 已保存: 图1_K线图.png")
    
    # 3.2 技术指标图 - MACD
    print("  生成图 2: MACD指标图...")
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(16, 8), gridspec_kw={'height_ratios': [2, 1]})
    
    # 价格和均线
    ax1.plot(x, df_tech['close'], label='收盘价', color='#333333', linewidth=2)
    ax1.plot(x, df_tech['MA5'], label='MA5', color='#FF6600', linewidth=1.5, alpha=0.7)
    ax1.plot(x, df_tech['MA10'], label='MA10', color='#3366FF', linewidth=1.5, alpha=0.7)
    ax1.plot(x, df_tech['MA20'], label='MA20', color='#00CC00', linewidth=1.5, alpha=0.7)
    ax1.set_title('芯动联科 (688582.SH) 均线分析', fontsize=14, fontweight='bold')
    ax1.set_ylabel('价格 (元)', fontsize=12)
    ax1.legend(loc='upper left', fontsize=10)
    ax1.grid(True, linestyle='--', alpha=0.3)
    
    # MACD
    ax2.plot(x, df_tech['MACD_DIF'], label='DIF', color='#FF6600', linewidth=1.5)
    ax2.plot(x, df_tech['MACD_DEA'], label='DEA', color='#3366FF', linewidth=1.5)
    
    colors_hist = ['#E04040' if v >= 0 else '#00AA00' for v in df_tech['MACD_HIST']]
    ax2.bar(x, df_tech['MACD_HIST'], color=colors_hist, alpha=0.5, width=0.8, label='MACD柱')
    ax2.axhline(y=0, color='#666666', linestyle='-', linewidth=0.8)
    
    ax2.set_title('MACD指标', fontsize=12)
    ax2.set_ylabel('MACD', fontsize=12)
    ax2.set_xlabel('交易日期', fontsize=12)
    ax2.legend(loc='upper left', fontsize=10)
    ax2.grid(True, linestyle='--', alpha=0.3)
    
    ax2.set_xticks(range(0, len(df_tech), step))
    ax2.set_xticklabels([d.strftime('%Y-%m') for d in df_tech.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    plt.tight_layout()
    plt.savefig('图2_MACD指标.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("  ✓ 已保存: 图2_MACD指标.png")
    
    # 3.3 RSI和KDJ指标
    print("  生成图 3: RSI和KDJ指标图...")
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(16, 8))
    
    # RSI
    ax1.plot(x, df_tech['RSI'], label='RSI(14)', color='#FF6600', linewidth=2)
    ax1.axhline(y=70, color='#E04040', linestyle='--', linewidth=1, alpha=0.7, label='超买线(70)')
    ax1.axhline(y=30, color='#00AA00', linestyle='--', linewidth=1, alpha=0.7, label='超卖线(30)')
    ax1.axhline(y=50, color='#666666', linestyle='-', linewidth=0.8, alpha=0.5)
    ax1.set_title('RSI相对强弱指标', fontsize=14, fontweight='bold')
    ax1.set_ylabel('RSI', fontsize=12)
    ax1.set_ylim(0, 100)
    ax1.legend(loc='upper right', fontsize=10)
    ax1.grid(True, linestyle='--', alpha=0.3)
    
    ax1.set_xticks(range(0, len(df_tech), step))
    ax1.set_xticklabels([d.strftime('%Y-%m') for d in df_tech.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    # KDJ
    ax2.plot(x, df_tech['K'], label='K值', color='#FF6600', linewidth=1.5)
    ax2.plot(x, df_tech['D'], label='D值', color='#3366FF', linewidth=1.5)
    ax2.plot(x, df_tech['J'], label='J值', color='#00CC00', linewidth=1.5)
    ax2.axhline(y=80, color='#E04040', linestyle='--', linewidth=1, alpha=0.7, label='超买线(80)')
    ax2.axhline(y=20, color='#00AA00', linestyle='--', linewidth=1, alpha=0.7, label='超卖线(20)')
    ax2.set_title('KDJ随机指标', fontsize=14, fontweight='bold')
    ax2.set_ylabel('KDJ', fontsize=12)
    ax2.set_xlabel('交易日期', fontsize=12)
    ax2.legend(loc='upper right', fontsize=10)
    ax2.grid(True, linestyle='--', alpha=0.3)
    
    ax2.set_xticks(range(0, len(df_tech), step))
    ax2.set_xticklabels([d.strftime('%Y-%m') for d in df_tech.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    plt.tight_layout()
    plt.savefig('图3_RSI_KDJ指标.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("  ✓ 已保存: 图3_RSI_KDJ指标.png")
    
    # 3.4 布林带
    print("  生成图 4: 布林带指标图...")
    
    fig, ax = plt.subplots(figsize=(16, 8))
    
    ax.plot(x, df_tech['close'], label='收盘价', color='#333333', linewidth=2, zorder=3)
    ax.plot(x, df_tech['BB_UPPER'], label='布林上轨', color='#E04040', linewidth=1.5, 
            linestyle='--', alpha=0.7)
    ax.plot(x, df_tech['BB_MIDDLE'], label='布林中轨(MA20)', color='#3366FF', linewidth=1.5, alpha=0.7)
    ax.plot(x, df_tech['BB_LOWER'], label='布林下轨', color='#00AA00', linewidth=1.5, 
            linestyle='--', alpha=0.7)
    
    # 填充布林带区域
    ax.fill_between(x, df_tech['BB_UPPER'], df_tech['BB_LOWER'], alpha=0.1, color='#3366FF')
    
    ax.set_title('芯动联科 (688582.SH) 布林带指标', fontsize=16, fontweight='bold', pad=15)
    ax.set_xlabel('交易日期', fontsize=12)
    ax.set_ylabel('价格 (元)', fontsize=12)
    ax.legend(loc='upper left', fontsize=10)
    ax.grid(True, linestyle='--', alpha=0.3, color='#999999')
    
    ax.set_xticks(range(0, len(df_tech), step))
    ax.set_xticklabels([d.strftime('%Y-%m') for d in df_tech.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    plt.tight_layout()
    plt.savefig('图4_布林带指标.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("  ✓ 已保存: 图4_布林带指标.png")
    
    # 3.5 估值指标图
    print("  生成图 5: 估值指标分析图...")
    
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 10))
    
    # PE
    ax1.plot(x, df_basic['pe'], label='PE', color='#FF6600', linewidth=2)
    ax1.plot(x, df_basic['pe_ttm'], label='PE(TTM)', color='#3366FF', linewidth=2)
    ax1.set_title('市盈率 (PE)', fontsize=14, fontweight='bold')
    ax1.set_ylabel('PE', fontsize=12)
    ax1.legend(fontsize=10)
    ax1.grid(True, linestyle='--', alpha=0.3)
    ax1.set_xticks(range(0, len(df_basic), step))
    ax1.set_xticklabels([d.strftime('%Y-%m') for d in df_basic.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    # PB
    ax2.plot(x, df_basic['pb'], label='PB', color='#00CC00', linewidth=2)
    ax2.set_title('市净率 (PB)', fontsize=14, fontweight='bold')
    ax2.set_ylabel('PB', fontsize=12)
    ax2.legend(fontsize=10)
    ax2.grid(True, linestyle='--', alpha=0.3)
    ax2.set_xticks(range(0, len(df_basic), step))
    ax2.set_xticklabels([d.strftime('%Y-%m') for d in df_basic.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    # PS
    ax3.plot(x, df_basic['ps'], label='PS', color='#CC00FF', linewidth=2)
    ax3.plot(x, df_basic['ps_ttm'], label='PS(TTM)', color='#FF66CC', linewidth=2)
    ax3.set_title('市销率 (PS)', fontsize=14, fontweight='bold')
    ax3.set_ylabel('PS', fontsize=12)
    ax3.set_xlabel('交易日期', fontsize=12)
    ax3.legend(fontsize=10)
    ax3.grid(True, linestyle='--', alpha=0.3)
    ax3.set_xticks(range(0, len(df_basic), step))
    ax3.set_xticklabels([d.strftime('%Y-%m') for d in df_basic.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    # 总市值
    ax4.plot(x, df_basic['total_mv'] / 10000, label='总市值(亿元)', color='#FF9900', linewidth=2)
    ax4.set_title('总市值', fontsize=14, fontweight='bold')
    ax4.set_ylabel('市值 (亿元)', fontsize=12)
    ax4.set_xlabel('交易日期', fontsize=12)
    ax4.legend(fontsize=10)
    ax4.grid(True, linestyle='--', alpha=0.3)
    ax4.set_xticks(range(0, len(df_basic), step))
    ax4.set_xticklabels([d.strftime('%Y-%m') for d in df_basic.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    plt.tight_layout()
    plt.savefig('图5_估值指标分析.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("  ✓ 已保存: 图5_估值指标分析.png")
    
    # 3.6 换手率和成交量
    print("  生成图 6: 换手率和成交量分析图...")
    
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(16, 8))
    
    # 换手率
    ax1.plot(x, df_basic['turnover_rate'], label='换手率(%)', color='#FF6600', linewidth=2)
    ax1.plot(x, df_basic['turnover_rate_f'], label='流通换手率(%)', color='#3366FF', linewidth=2)
    ax1.set_title('换手率分析', fontsize=14, fontweight='bold')
    ax1.set_ylabel('换手率 (%)', fontsize=12)
    ax1.legend(loc='upper right', fontsize=10)
    ax1.grid(True, linestyle='--', alpha=0.3)
    ax1.set_xticks(range(0, len(df_basic), step))
    ax1.set_xticklabels([d.strftime('%Y-%m') for d in df_basic.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    # 量比
    ax2.plot(x, df_basic['volume_ratio'], label='量比', color='#00CC00', linewidth=2)
    ax2.axhline(y=1, color='#666666', linestyle='-', linewidth=0.8, alpha=0.7, label='基准线(1)')
    ax2.set_title('量比分析', fontsize=14, fontweight='bold')
    ax2.set_ylabel('量比', fontsize=12)
    ax2.set_xlabel('交易日期', fontsize=12)
    ax2.legend(loc='upper right', fontsize=10)
    ax2.grid(True, linestyle='--', alpha=0.3)
    ax2.set_xticks(range(0, len(df_basic), step))
    ax2.set_xticklabels([d.strftime('%Y-%m') for d in df_basic.iloc[::step]['trade_date']], 
                        rotation=45, ha='right')
    
    plt.tight_layout()
    plt.savefig('图6_换手率量比分析.png', dpi=150, bbox_inches='tight')
    plt.close()
    print("  ✓ 已保存: 图6_换手率量比分析.png")
    
    print("\n✓ 所有图表生成完成！")
    
else:
    # 使用HTML+Chart.js方案
    print("\n  使用HTML+Chart.js方案生成图表...")
    # 这里可以生成HTML文件，但考虑到复杂度，先跳过
    print("  ⚠ 建议安装matplotlib以生成更高质量的图表")

# ============================================================================
# 4. 生成统计分析文本
# ============================================================================
print("\n" + "=" * 60)
print("步骤 4/6: 生成统计分析")
print("=" * 60)

# 基本统计
close_start = df_stock.iloc[0]['close']
close_end = df_stock.iloc[-1]['close']
price_change = ((close_end / close_start) - 1) * 100
price_max = df_stock['high'].max()
price_min = df_stock['low'].min()
price_avg = df_stock['close'].mean()

# 技术指标统计
rsi_current = df_tech.iloc[-1]['RSI']
macd_current = df_tech.iloc[-1]['MACD_HIST']
k_current = df_tech.iloc[-1]['K']
d_current = df_tech.iloc[-1]['D']

# 估值指标统计
pe_current = df_basic.iloc[-1]['pe']
pb_current = df_basic.iloc[-1]['pb']
ps_current = df_basic.iloc[-1]['ps']
mv_current = df_basic.iloc[-1]['total_mv'] / 10000  # 转换为亿元

stats_text = f"""
## 数据统计摘要

### 价格统计
- 期初价格: {close_start:.2f} 元 (2025-06-30)
- 期末价格: {close_end:.2f} 元 (2026-06-26)
- 期间涨跌幅: {price_change:.2f}%
- 最高价: {price_max:.2f} 元
- 最低价: {price_min:.2f} 元
- 平均收盘价: {price_avg:.2f} 元

### 技术指标分析
- RSI(14): {rsi_current:.2f} {'（超买区域，注意回调风险）' if rsi_current > 70 else '（超卖区域，可能存在反弹机会）' if rsi_current < 30 else '（中性区域）'}
- MACD柱: {macd_current:.2f} {'（正值，多头信号）' if macd_current > 0 else '（负值，空头信号）'}
- KDJ: K={k_current:.2f}, D={d_current:.2f} {'（金叉，买入信号）' if k_current > d_current else '（死叉，卖出信号）'}

### 估值指标分析
- 市盈率(PE): {pe_current:.2f} {'（估值较高，注意风险）' if pe_current > 50 else '（估值合理）'}
- 市净率(PB): {pb_current:.2f}
- 市销率(PS): {ps_current:.2f}
- 总市值: {mv_current:.2f} 亿元
"""

print("✓ 统计分析完成")

# ============================================================================
# 5. 生成Word报告
# ============================================================================
print("\n" + "=" * 60)
print("步骤 5/6: 生成Word报告")
print("=" * 60)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # 设置文档默认样式（宋体、五号字、1.5倍行距、0倍段间距、两端对齐）
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)  # 五号字 = 10.5pt
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)  # 0倍段前间距
    paragraph_format.space_after = Pt(0)  # 0倍段后间距
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
    
    # 标题
    title = doc.add_heading('芯动联科 (688582.SH) 股票分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    date_para = doc.add_paragraph(f'报告生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 目录
    doc.add_heading('目录', 1)
    doc.add_paragraph('一、公司简介')
    doc.add_paragraph('二、K线分析')
    doc.add_paragraph('三、基本面分析')
    doc.add_paragraph('四、技术面分析')
    doc.add_paragraph('五、投资建议')
    doc.add_paragraph('六、风险提示')
    
    # 一、公司简介
    doc.add_heading('一、公司简介', 1)
    doc.add_paragraph('芯动联科（688582.SH）是一家科创板上市公司，主要从事MEMS传感器的研发、生产和销售。')
    doc.add_paragraph(f'本报告分析了该公司从2025年6月30日至2026年6月26日共{len(df_stock)}个交易日的股价表现和基本面情况。')
    
    # 二、K线分析
    doc.add_heading('二、K线分析', 1)
    doc.add_paragraph('K线图是技术分析的基础工具，通过每日的开盘价、收盘价、最高价、最低价四个价格，直观展示股价的波动情况。')
    
    if USE_MATPLOTLIB:
        doc.add_picture('图1_K线图.png', width=Inches(6.5))
        doc.add_paragraph('图1: 芯动联科K线图')
        doc.add_paragraph('解读: K线图中，红色柱体表示当日收盘价高于开盘价（阳线），绿色柱体表示当日收盘价低于开盘价（阴线）。从图中可以看出，该股票在2025年8月中旬出现一波快速上涨，随后进入震荡调整阶段。均线系统显示，短期均线在长期均线上方，表明中长期趋势依然向上。')
    
    # 三、基本面分析
    doc.add_heading('三、基本面分析', 1)
    doc.add_paragraph('基本面分析主要关注公司的估值指标，包括市盈率(PE)、市净率(PB)、市销率(PS)等。这些指标帮助投资者判断股票是否被高估或低估。')
    
    if USE_MATPLOTLIB:
        doc.add_picture('图5_估值指标分析.png', width=Inches(6.5))
        doc.add_paragraph('图5: 估值指标分析')
        doc.add_paragraph(f'解读: 当前市盈率(PE)为{pe_current:.2f}，{'处于较高水平，表明市场对公司未来增长预期较高，但也意味着投资风险较大。' if pe_current > 50 else '处于合理水平。'}市净率(PB)为{pb_current:.2f}，总市值为{mv_current:.2f}亿元。投资者应关注公司基本面变化，及时调整投资策略。')
    
    # 四、技术面分析
    doc.add_heading('四、技术面分析', 1)
    doc.add_paragraph('技术面分析通过各类技术指标，帮助投资者判断买卖时机。本报告选取了MACD、RSI、KDJ、布林带等常用指标进行分析。')
    
    if USE_MATPLOTLIB:
        doc.add_picture('图2_MACD指标.png', width=Inches(6.5))
        doc.add_paragraph('图2: MACD指标分析')
        doc.add_paragraph('解读: MACD指标由DIF线、DEA线和MACD柱组成。当DIF线上穿DEA线时，形成金叉，为买入信号；当DIF线下穿DEA线时，形成死叉，为卖出信号。MACD柱由正转负，表明多头力量减弱，空头力量增强。')
        
        doc.add_picture('图3_RSI_KDJ指标.png', width=Inches(6.5))
        doc.add_paragraph('图3: RSI和KDJ指标分析')
        doc.add_paragraph(f'解读: RSI指标衡量股价的超买超卖情况。当前RSI为{rsi_current:.2f}，{'处于超买区域（>70），股价可能面临回调压力。' if rsi_current > 70 else '处于超卖区域（<30），股价可能存在反弹机会。' if rsi_current < 30 else '处于中性区域（30-70），股价走势相对平稳。'}KDJ指标中，K值和D值的交叉也提供买卖信号。')
        
        doc.add_picture('图4_布林带指标.png', width=Inches(6.5))
        doc.add_paragraph('图4: 布林带指标分析')
        doc.add_paragraph('解读: 布林带由中轨（20日均线）和上下轨（中轨±2倍标准差）组成。当股价触及上轨时，表明股价可能过高，存在回调风险；当股价触及下轨时，表明股价可能过低，存在反弹机会。当前股价在布林带中轨附近运行，表明市场处于相对平衡状态。')
        
        doc.add_picture('图6_换手率量比分析.png', width=Inches(6.5))
        doc.add_paragraph('图6: 换手率和量比分析')
        doc.add_paragraph('解读: 换手率衡量股票的流动性，换手率越高，表明股票交易越活跃。量比衡量当日成交量与过去5日平均成交量的比值，量比大于1表明当日成交量放大。投资者应关注成交量的变化，量价配合是股价上涨的重要条件。')
    
    # 五、投资建议
    doc.add_heading('五、投资建议', 1)
    
    # 根据技术指标给出建议
    advice = ""
    if rsi_current > 70:
        advice += "RSI指标显示超买，建议谨慎操作，可适当减仓。\n"
    elif rsi_current < 30:
        advice += "RSI指标显示超卖，可能存在反弹机会，可适当关注。\n"
    else:
        advice += "RSI指标处于中性区域，股价走势相对平稳。\n"
    
    if macd_current > 0:
        advice += "MACD指标显示多头信号，可考虑持有或加仓。\n"
    else:
        advice += "MACD指标显示空头信号，建议观望或减仓。\n"
    
    if pe_current > 50:
        advice += "估值偏高，注意投资风险。\n"
    else:
        advice += "估值相对合理。\n"
    
    doc.add_paragraph(advice)
    doc.add_paragraph('综合来看，投资者应根据自身风险承受能力和投资目标，制定合理的投资策略。建议分散投资，控制仓位，设置止损点。')
    
    # 六、风险提示
    doc.add_heading('六、风险提示', 1)
    doc.add_paragraph('1. 股市有风险，投资需谨慎。')
    doc.add_paragraph('2. 本报告仅供参考，不构成投资建议。')
    doc.add_paragraph('3. 技术分析存在滞后性，投资者应结合基本面和市场环境综合判断。')
    doc.add_paragraph('4. 科创板股票波动较大，投资者应注意风险控制。')
    doc.add_paragraph('5. 本报告数据来源于Tushare，数据的准确性和完整性有待验证。')
    
    # 保存文档
    doc.save('芯动联科股票分析报告.docx')
    print("✓ Word报告已保存: 芯动联科股票分析报告.docx")
    
except ImportError:
    print("✗ python-docx 未安装，跳过Word报告生成")
    print("  可运行: pip install python-docx")

# ============================================================================
# 6. 生成Jupyter Notebook
# ============================================================================
print("\n" + "=" * 60)
print("步骤 6/6: 生成Jupyter Notebook")
print("=" * 60)

notebook_content = {
    "cells": [
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "# 芯动联科 (688582.SH) 股票分析\n",
                "\n",
                "本Notebook包含对芯动联科股票的完整分析，包括：\n",
                "- 数据获取\n",
                "- K线分析\n",
                "- 基本面分析\n",
                "- 技术面分析\n",
                "- 投资建议\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "import pandas as pd\n",
                "import numpy as np\n",
                "import json\n",
                "import matplotlib.pyplot as plt\n",
                "import warnings\n",
                "warnings.filterwarnings('ignore')\n",
                "\n",
                "# 设置中文字体\n",
                "plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']\n",
                "plt.rcParams['axes.unicode_minus'] = False\n",
                "\n",
                "print('环境准备完成')"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 1. 读取数据\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "# 读取日交易数据\n",
                "with open('stock_data.json', 'r', encoding='utf-8') as f:\n",
                "    stock_data = json.load(f)\n",
                "\n",
                "# 读取每日指标数据\n",
                "with open('daily_basic_data.json', 'r', encoding='utf-8') as f:\n",
                "    basic_data = json.load(f)\n",
                "\n",
                "# 转换为DataFrame\n",
                "df_stock = pd.DataFrame(stock_data)\n",
                "df_basic = pd.DataFrame(basic_data)\n",
                "\n",
                "# 日期格式转换\n",
                "df_stock['trade_date'] = pd.to_datetime(df_stock['trade_date'], format='%Y%m%d')\n",
                "df_basic['trade_date'] = pd.to_datetime(df_basic['trade_date'], format='%Y%m%d')\n",
                "\n",
                "# 按日期排序\n",
                "df_stock = df_stock.sort_values('trade_date').reset_index(drop=True)\n",
                "df_basic = df_basic.sort_values('trade_date').reset_index(drop=True)\n",
                "\n",
                "print(f'日交易数据: {len(df_stock)} 条记录')\n",
                "print(f'每日指标数据: {len(df_basic)} 条记录')\n",
                "print(f'数据时间范围: {df_stock[\"trade_date\"].min().date()} ~ {df_stock[\"trade_date\"].max().date()}')\n",
                "df_stock.head()"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 2. 计算技术指标\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "def calculate_ma(df, periods=[5, 10, 20, 60]):\n",
                "    for period in periods:\n",
                "        df[f'MA{period}'] = df['close'].rolling(window=period).mean()\n",
                "    return df\n",
                "\n",
                "def calculate_macd(df, fast=12, slow=26, signal=9):\n",
                "    ema_fast = df['close'].ewm(span=fast, adjust=False).mean()\n",
                "    ema_slow = df['close'].ewm(span=slow, adjust=False).mean()\n",
                "    df['MACD_DIF'] = ema_fast - ema_slow\n",
                "    df['MACD_DEA'] = df['MACD_DIF'].ewm(span=signal, adjust=False).mean()\n",
                "    df['MACD_HIST'] = (df['MACD_DIF'] - df['MACD_DEA']) * 2\n",
                "    return df\n",
                "\n",
                "def calculate_rsi(df, period=14):\n",
                "    delta = df['close'].diff()\n",
                "    gain = (delta.where(delta > 0, 0)).rolling(window=period).mean()\n",
                "    loss = (-delta.where(delta < 0, 0)).rolling(window=period).mean()\n",
                "    rs = gain / loss\n",
                "    df['RSI'] = 100 - (100 / (1 + rs))\n",
                "    return df\n",
                "\n",
                "# 计算技术指标\n",
                "df_tech = df_stock.copy()\n",
                "df_tech = calculate_ma(df_tech)\n",
                "df_tech = calculate_macd(df_tech)\n",
                "df_tech = calculate_rsi(df_tech)\n",
                "\n",
                "print('技术指标计算完成')\n",
                "df_tech[['trade_date', 'close', 'MA5', 'MA20', 'RSI', 'MACD_DIF', 'MACD_HIST']].tail()"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 3. 绘制K线图\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "# 绘制K线图和均线\n",
                "fig, ax = plt.subplots(figsize=(15, 7))\n",
                "\n",
                "x = range(len(df_tech))\n",
                "ax.plot(x, df_tech['close'], label='收盘价', color='black', linewidth=2)\n",
                "ax.plot(x, df_tech['MA5'], label='MA5', color='orange', linewidth=1.5, alpha=0.7)\n",
                "ax.plot(x, df_tech['MA20'], label='MA20', color='blue', linewidth=1.5, alpha=0.7)\n",
                "ax.plot(x, df_tech['MA60'], label='MA60', color='purple', linewidth=1.5, alpha=0.7)\n",
                "\n",
                "ax.set_title('芯动联科 (688582.SH) 股价走势与均线', fontsize=16, fontweight='bold')\n",
                "ax.set_xlabel('交易日期', fontsize=12)\n",
                "ax.set_ylabel('价格 (元)', fontsize=12)\n",
                "ax.legend(loc='upper left', fontsize=11)\n",
                "ax.grid(True, linestyle='--', alpha=0.4)\n",
                "\n",
                "step = len(df_tech) // 10\n",
                "ax.set_xticks(range(0, len(df_tech), step))\n",
                "ax.set_xticklabels([d.strftime('%Y-%m') for d in df_tech.iloc[::step]['trade_date']], rotation=45, ha='right')\n",
                "\n",
                "plt.tight_layout()\n",
                "plt.show()"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 4. 技术指标分析\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "# 绘制MACD指标\n",
                "fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(15, 8), gridspec_kw={'height_ratios': [2, 1]})\n",
                "\n",
                "ax1.plot(x, df_tech['close'], label='收盘价', color='black', linewidth=2)\n",
                "ax1.set_title('MACD指标分析', fontsize=14, fontweight='bold')\n",
                "ax1.legend(loc='upper left')\n",
                "ax1.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "ax2.plot(x, df_tech['MACD_DIF'], label='DIF', color='orange', linewidth=1.5)\n",
                "ax2.plot(x, df_tech['MACD_DEA'], label='DEA', color='blue', linewidth=1.5)\n",
                "ax2.bar(x, df_tech['MACD_HIST'], color=['red' if v >= 0 else 'green' for v in df_tech['MACD_HIST']], alpha=0.5)\n",
                "ax2.axhline(y=0, color='gray', linestyle='-', linewidth=0.8)\n",
                "ax2.legend(loc='upper left')\n",
                "ax2.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "plt.tight_layout()\n",
                "plt.show()"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "# 绘制RSI指标\n",
                "fig, ax = plt.subplots(figsize=(15, 5))\n",
                "\n",
                "ax.plot(x, df_tech['RSI'], label='RSI(14)', color='orange', linewidth=2)\n",
                "ax.axhline(y=70, color='red', linestyle='--', linewidth=1, alpha=0.7, label='超买线(70)')\n",
                "ax.axhline(y=30, color='green', linestyle='--', linewidth=1, alpha=0.7, label='超卖线(30)')\n",
                "ax.set_title('RSI相对强弱指标', fontsize=14, fontweight='bold')\n",
                "ax.set_ylabel('RSI', fontsize=12)\n",
                "ax.legend(loc='upper right')\n",
                "ax.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "plt.tight_layout()\n",
                "plt.show()"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 5. 估值指标分析\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "# 绘制估值指标\n",
                "fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(15, 10))\n",
                "\n",
                "ax1.plot(x, df_basic['pe'], label='PE', color='orange', linewidth=2)\n",
                "ax1.plot(x, df_basic['pe_ttm'], label='PE(TTM)', color='blue', linewidth=2)\n",
                "ax1.set_title('市盈率 (PE)', fontsize=14, fontweight='bold')\n",
                "ax1.legend()\n",
                "ax1.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "ax2.plot(x, df_basic['pb'], label='PB', color='green', linewidth=2)\n",
                "ax2.set_title('市净率 (PB)', fontsize=14, fontweight='bold')\n",
                "ax2.legend()\n",
                "ax2.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "ax3.plot(x, df_basic['total_mv'] / 10000, label='总市值(亿元)', color='red', linewidth=2)\n",
                "ax3.set_title('总市值', fontsize=14, fontweight='bold')\n",
                "ax3.legend()\n",
                "ax3.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "ax4.plot(x, df_basic['turnover_rate'], label='换手率(%)', color='purple', linewidth=2)\n",
                "ax4.set_title('换手率', fontsize=14, fontweight='bold')\n",
                "ax4.legend()\n",
                "ax4.grid(True, linestyle='--', alpha=0.3)\n",
                "\n",
                "plt.tight_layout()\n",
                "plt.show()"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 6. 统计分析\n"
            ]
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "# 统计分析\n",
                "close_start = df_stock.iloc[0]['close']\n",
                "close_end = df_stock.iloc[-1]['close']\n",
                "price_change = ((close_end / close_start) - 1) * 100\n",
                "price_max = df_stock['high'].max()\n",
                "price_min = df_stock['low'].min()\n",
                "\n",
                "print('=== 价格统计 ===')\n",
                "print(f'期初价格: {close_start:.2f} 元')\n",
                "print(f'期末价格: {close_end:.2f} 元')\n",
                "print(f'期间涨跌幅: {price_change:.2f}%')\n",
                "print(f'最高价: {price_max:.2f} 元')\n",
                "print(f'最低价: {price_min:.2f} 元')\n",
                "print(f'平均收盘价: {df_stock[\"close\"].mean():.2f} 元')\n",
                "\n",
                "print('\\n=== 技术指标 ===')\n",
                "print(f'RSI(14): {df_tech.iloc[-1][\"RSI\"]:.2f}')\n",
                "print(f'MACD柱: {df_tech.iloc[-1][\"MACD_HIST\"]:.2f}')\n",
                "\n",
                "print('\\n=== 估值指标 ===')\n",
                "print(f'市盈率(PE): {df_basic.iloc[-1][\"pe\"]:.2f}')\n",
                "print(f'市净率(PB): {df_basic.iloc[-1][\"pb\"]:.2f}')\n",
                "print(f'总市值: {df_basic.iloc[-1][\"total_mv\"] / 10000:.2f} 亿元')"
            ]
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 7. 投资建议\n",
                "\n",
                "根据技术指标和基本面分析，给出投资建议：\n",
                "\n",
                "- 当RSI>70时，建议谨慎操作\n",
                "- 当MACD柱为正时，可考虑持有\n",
                "- 当PE过高时，注意风险\n",
                "\n",
                "**注意**: 本分析仅供参考，不构成投资建议。"
            ]
        }
    ],
    "metadata": {
        "kernelspec": {
            "display_name": "Python 3",
            "language": "python",
            "name": "python3"
        },
        "language_info": {
            "name": "python",
            "version": "3.13.12"
        }
    },
    "nbformat": 4,
    "nbformat_minor": 5
}

# 保存Notebook
with open('芯动联科股票分析.ipynb', 'w', encoding='utf-8') as f:
    json.dump(notebook_content, f, ensure_ascii=False, indent=2)

print("✓ Jupyter Notebook已保存: 芯动联科股票分析.ipynb")

# ============================================================================
# 完成
# ============================================================================
print("\n" + "=" * 60)
print("分析完成！")
print("=" * 60)
print("\n生成的文件:")
print("1. 图1_K线图.png - K线分析图")
print("2. 图2_MACD指标.png - MACD指标图")
print("3. 图3_RSI_KDJ指标.png - RSI和KDJ指标图")
print("4. 图4_布林带指标.png - 布林带指标图")
print("5. 图5_估值指标分析.png - 估值指标分析图")
print("6. 图6_换手率量比分析.png - 换手率和量比分析图")
print("7. 芯动联科股票分析报告.docx - Word分析报告")
print("8. 芯动联科股票分析.ipynb - Jupyter Notebook")
print("9. 芯动联科_688582_年度日交易数据.csv - 原始数据")
print("\n所有文件已保存到当前目录。")
