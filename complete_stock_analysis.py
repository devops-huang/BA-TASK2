#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
芯动联科（688582.SH）股票完整分析脚本
生成6个技术指标图表和完整Word报告
"""

import json
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimSun', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

class StockAnalyzer:
    def __init__(self, stock_data_file, basic_data_file):
        """初始化股票分析器"""
        # 读取数据
        with open(stock_data_file, 'r', encoding='utf-8') as f:
            self.stock_data = json.load(f)
        with open(basic_data_file, 'r', encoding='utf-8') as f:
            self.basic_data = json.load(f)

        # 转换为DataFrame
        self.df_stock = pd.DataFrame(self.stock_data)
        self.df_basic = pd.DataFrame(self.basic_data)

        # 数据处理
        self.df_stock['trade_date'] = pd.to_datetime(self.df_stock['trade_date'], format='%Y%m%d')
        self.df_basic['trade_date'] = pd.to_datetime(self.df_basic['trade_date'], format='%Y%m%d')

        # 按日期排序
        self.df_stock = self.df_stock.sort_values('trade_date').reset_index(drop=True)
        self.df_basic = self.df_basic.sort_values('trade_date').reset_index(drop=True)

        # 合并数据（只选择需要的列，避免列名冲突）
        basic_cols = ['trade_date', 'pe', 'pe_ttm', 'pb', 'ps', 'ps_ttm',
                      'total_share', 'float_share', 'free_share',
                      'total_mv', 'circ_mv', 'turnover_rate',
                      'turnover_rate_f', 'volume_ratio']
        self.df = pd.merge(self.df_stock, self.df_basic[basic_cols], on='trade_date', how='left')

        print(f"数据加载完成：共{len(self.df)}条交易记录")

    def calculate_ma(self):
        """计算移动平均线"""
        self.df['MA5'] = self.df['close'].rolling(window=5).mean()
        self.df['MA10'] = self.df['close'].rolling(window=10).mean()
        self.df['MA20'] = self.df['close'].rolling(window=20).mean()
        self.df['MA60'] = self.df['close'].rolling(window=60).mean()

    def calculate_macd(self):
        """计算MACD指标"""
        exp1 = self.df['close'].ewm(span=12, adjust=False).mean()
        exp2 = self.df['close'].ewm(span=26, adjust=False).mean()
        self.df['MACD_DIF'] = exp1 - exp2
        self.df['MACD_DEA'] = self.df['MACD_DIF'].ewm(span=9, adjust=False).mean()
        self.df['MACD_HIST'] = (self.df['MACD_DIF'] - self.df['MACD_DEA']) * 2

    def calculate_rsi(self, period=14):
        """计算RSI指标"""
        delta = self.df['close'].diff()
        gain = (delta.where(delta > 0, 0)).rolling(window=period).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(window=period).mean()
        rs = gain / loss
        self.df['RSI'] = 100 - (100 / (1 + rs))

    def calculate_kdj(self):
        """计算KDJ指标"""
        low_list = self.df['low'].rolling(window=9).min()
        high_list = self.df['high'].rolling(window=9).max()
        rsv = (self.df['close'] - low_list) / (high_list - low_list) * 100

        self.df['K'] = rsv.ewm(com=2, adjust=False).mean()
        self.df['D'] = self.df['K'].ewm(com=2, adjust=False).mean()
        self.df['J'] = 3 * self.df['K'] - 2 * self.df['D']

    def calculate_bollinger_bands(self, period=20, std_dev=2):
        """计算布林带指标"""
        self.df['BB_MIDDLE'] = self.df['close'].rolling(window=period).mean()
        std = self.df['close'].rolling(window=period).std()
        self.df['BB_UPPER'] = self.df['BB_MIDDLE'] + (std * std_dev)
        self.df['BB_LOWER'] = self.df['BB_MIDDLE'] - (std * std_dev)

    def generate_all_indicators(self):
        """生成所有技术指标"""
        print("正在计算技术指标...")
        self.calculate_ma()
        self.calculate_macd()
        self.calculate_rsi()
        self.calculate_kdj()
        self.calculate_bollinger_bands()
        print("技术指标计算完成！")

    def plot_kline_with_ma(self):
        """图1: K线图与均线分析"""
        fig, ax = plt.subplots(figsize=(14, 8))
        fig.patch.set_facecolor('white')

        # 绘制K线
        for i in range(len(self.df)):
            date = self.df.iloc[i]['trade_date']
            open_price = self.df.iloc[i]['open']
            close_price = self.df.iloc[i]['close']
            high_price = self.df.iloc[i]['high']
            low_price = self.df.iloc[i]['low']

            color = 'red' if close_price >= open_price else 'green'
            alpha = 0.8

            # 实体
            ax.plot([date, date], [low_price, high_price], color=color, linewidth=1, alpha=alpha)
            ax.plot([date, date], [open_price, close_price], color=color, linewidth=4, alpha=alpha)

        # 绘制均线
        ax.plot(self.df['trade_date'], self.df['MA5'], label='MA5', linewidth=1.5, alpha=0.8)
        ax.plot(self.df['trade_date'], self.df['MA10'], label='MA10', linewidth=1.5, alpha=0.8)
        ax.plot(self.df['trade_date'], self.df['MA20'], label='MA20', linewidth=1.5, alpha=0.8)
        ax.plot(self.df['trade_date'], self.df['MA60'], label='MA60', linewidth=1.5, alpha=0.8)

        ax.set_title('芯动联科（688582.SH）- K线图与均线分析', fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('日期', fontsize=12)
        ax.set_ylabel('价格（元）', fontsize=12)
        ax.legend(loc='best', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('图1_K线图与均线分析.png', dpi=150, bbox_inches='tight')
        plt.close()
        print("✓ 图1生成完成：K线图与均线分析")

    def plot_macd(self):
        """图2: MACD指标分析"""
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10), height_ratios=[2, 1])
        fig.patch.set_facecolor('white')

        # 上图：收盘价
        ax1.plot(self.df['trade_date'], self.df['close'], color='black', linewidth=2, label='收盘价')
        ax1.set_ylabel('价格（元）', fontsize=12)
        ax1.set_title('MACD指标分析', fontsize=16, fontweight='bold', pad=20)
        ax1.legend(loc='best')
        ax1.grid(True, alpha=0.3)

        # 下图：MACD
        ax2.plot(self.df['trade_date'], self.df['MACD_DIF'], label='DIF', linewidth=1.5)
        ax2.plot(self.df['trade_date'], self.df['MACD_DEA'], label='DEA', linewidth=1.5)

        # MACD柱状图
        colors = ['red' if x > 0 else 'green' for x in self.df['MACD_HIST']]
        ax2.bar(self.df['trade_date'], self.df['MACD_HIST'], color=colors, alpha=0.5, label='MACD')

        ax2.set_xlabel('日期', fontsize=12)
        ax2.set_ylabel('MACD', fontsize=12)
        ax2.legend(loc='best')
        ax2.grid(True, alpha=0.3)
        ax2.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))

        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('图2_MACD指标分析.png', dpi=150, bbox_inches='tight')
        plt.close()
        print("✓ 图2生成完成：MACD指标分析")

    def plot_rsi_kdj(self):
        """图3: RSI和KDJ指标分析"""
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10))
        fig.patch.set_facecolor('white')

        # RSI
        ax1.plot(self.df['trade_date'], self.df['RSI'], label='RSI(14)', color='purple', linewidth=2)
        ax1.axhline(y=70, color='red', linestyle='--', alpha=0.5, label='超买线(70)')
        ax1.axhline(y=30, color='green', linestyle='--', alpha=0.5, label='超卖线(30)')
        ax1.axhline(y=50, color='gray', linestyle='-', alpha=0.3)
        ax1.set_ylabel('RSI', fontsize=12)
        ax1.set_title('RSI和KDJ指标分析', fontsize=16, fontweight='bold', pad=20)
        ax1.legend(loc='best')
        ax1.grid(True, alpha=0.3)
        ax1.set_ylim(0, 100)

        # KDJ
        ax2.plot(self.df['trade_date'], self.df['K'], label='K值', linewidth=1.5)
        ax2.plot(self.df['trade_date'], self.df['D'], label='D值', linewidth=1.5)
        ax2.plot(self.df['trade_date'], self.df['J'], label='J值', linewidth=1.5)
        ax2.axhline(y=80, color='red', linestyle='--', alpha=0.5, label='超买线(80)')
        ax2.axhline(y=20, color='green', linestyle='--', alpha=0.5, label='超卖线(20)')
        ax2.axhline(y=50, color='gray', linestyle='-', alpha=0.3)
        ax2.set_xlabel('日期', fontsize=12)
        ax2.set_ylabel('KDJ', fontsize=12)
        ax2.legend(loc='best')
        ax2.grid(True, alpha=0.3)
        ax2.set_ylim(0, 100)
        ax2.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))

        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('图3_RSI和KDJ指标分析.png', dpi=150, bbox_inches='tight')
        plt.close()
        print("✓ 图3生成完成：RSI和KDJ指标分析")

    def plot_bollinger_bands(self):
        """图4: 布林带指标"""
        fig, ax = plt.subplots(figsize=(14, 8))
        fig.patch.set_facecolor('white')

        # K线简化显示
        ax.plot(self.df['trade_date'], self.df['close'], color='black', linewidth=2, label='收盘价', alpha=0.7)

        # 布林带
        ax.plot(self.df['trade_date'], self.df['BB_UPPER'], label='上轨', linewidth=1.5, linestyle='--', color='red')
        ax.plot(self.df['trade_date'], self.df['BB_MIDDLE'], label='中轨(MA20)', linewidth=1.5, color='blue')
        ax.plot(self.df['trade_date'], self.df['BB_LOWER'], label='下轨', linewidth=1.5, linestyle='--', color='green')

        # 填充区域
        ax.fill_between(self.df['trade_date'], self.df['BB_UPPER'], self.df['BB_LOWER'], alpha=0.1, color='gray')

        ax.set_title('芯动联科（688582.SH）- 布林带指标', fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('日期', fontsize=12)
        ax.set_ylabel('价格（元）', fontsize=12)
        ax.legend(loc='best', fontsize=10)
        ax.grid(True, alpha=0.3)
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('图4_布林带指标.png', dpi=150, bbox_inches='tight')
        plt.close()
        print("✓ 图4生成完成：布林带指标")

    def plot_valuation(self):
        """图5: 估值指标分析"""
        fig, axes = plt.subplots(2, 2, figsize=(16, 12))
        fig.patch.set_facecolor('white')
        fig.suptitle('估值指标分析', fontsize=16, fontweight='bold')

        # PE
        axes[0, 0].plot(self.df['trade_date'], self.df['pe'], color='blue', linewidth=2)
        axes[0, 0].set_title('市盈率 PE', fontsize=12)
        axes[0, 0].set_ylabel('PE')
        axes[0, 0].grid(True, alpha=0.3)

        # PB
        axes[0, 1].plot(self.df['trade_date'], self.df['pb'], color='green', linewidth=2)
        axes[0, 1].set_title('市净率 PB', fontsize=12)
        axes[0, 1].set_ylabel('PB')
        axes[0, 1].grid(True, alpha=0.3)

        # PS
        axes[1, 0].plot(self.df['trade_date'], self.df['ps'], color='red', linewidth=2)
        axes[1, 0].set_title('市销率 PS', fontsize=12)
        axes[1, 0].set_ylabel('PS')
        axes[1, 0].set_xlabel('日期')
        axes[1, 0].grid(True, alpha=0.3)

        # 总市值
        axes[1, 1].plot(self.df['trade_date'], self.df['total_mv'] / 10000, color='purple', linewidth=2)
        axes[1, 1].set_title('总市值（亿元）', fontsize=12)
        axes[1, 1].set_ylabel('市值（亿元）')
        axes[1, 1].set_xlabel('日期')
        axes[1, 1].grid(True, alpha=0.3)

        for ax in axes.flat:
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
            plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right')

        plt.tight_layout()
        plt.savefig('图5_估值指标分析.png', dpi=150, bbox_inches='tight')
        plt.close()
        print("✓ 图5生成完成：估值指标分析")

    def plot_turnover(self):
        """图6: 换手率和量比分析"""
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10))
        fig.patch.set_facecolor('white')

        # 换手率
        ax1.plot(self.df['trade_date'], self.df['turnover_rate'], label='换手率(%)', color='blue', linewidth=2)
        ax1.plot(self.df['trade_date'], self.df['turnover_rate_f'], label='换手率(自由流通)', color='cyan', linewidth=2, linestyle='--')
        ax1.set_ylabel('换手率(%)', fontsize=12)
        ax1.set_title('换手率和量比分析', fontsize=16, fontweight='bold', pad=20)
        ax1.legend(loc='best')
        ax1.grid(True, alpha=0.3)

        # 量比
        ax2.plot(self.df['trade_date'], self.df['volume_ratio'], label='量比', color='red', linewidth=2)
        ax2.axhline(y=1, color='gray', linestyle='-', alpha=0.5, label='基准线(1.0)')
        ax2.fill_between(self.df['trade_date'], 0.8, 1.2, alpha=0.2, color='gray', label='正常区间(0.8-1.2)')
        ax2.set_xlabel('日期', fontsize=12)
        ax2.set_ylabel('量比', fontsize=12)
        ax2.legend(loc='best')
        ax2.grid(True, alpha=0.3)
        ax2.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))

        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.savefig('图6_换手率和量比分析.png', dpi=150, bbox_inches='tight')
        plt.close()
        print("✓ 图6生成完成：换手率和量比分析")

    def generate_all_charts(self):
        """生成所有图表"""
        print("\n开始生成图表...")
        self.plot_kline_with_ma()
        self.plot_macd()
        self.plot_rsi_kdj()
        self.plot_bollinger_bands()
        self.plot_valuation()
        self.plot_turnover()
        print("\n✓ 所有图表生成完成！\n")

    def generate_word_report(self):
        """生成Word分析报告"""
        print("开始生成Word报告...")
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)  # 五号字
        style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        style.paragraph_format.space_before = Pt(0)  # 段前距0
        style.paragraph_format.space_after = Pt(0)   # 段后距0
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐

        # 封面
        doc.add_heading('芯动联科（688582.SH）股票分析报告', 0)
        doc.add_paragraph()
        doc.add_paragraph('完整版')
        doc.add_paragraph()
        doc.add_paragraph(f'报告生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
        doc.add_paragraph(f'数据期间：{self.df["trade_date"].min().strftime("%Y-%m-%d")} 至 {self.df["trade_date"].max().strftime("%Y-%m-%d")}')
        doc.add_page_break()

        # 目录
        doc.add_heading('目录', 1)
        doc.add_paragraph('1. 公司概况')
        doc.add_paragraph('2. K线分析')
        doc.add_paragraph('3. 基本面分析')
        doc.add_paragraph('4. 技术面分析')
        doc.add_paragraph('5. 投资建议')
        doc.add_paragraph('6. 风险提示')
        doc.add_page_break()

        # 1. 公司概况
        doc.add_heading('1. 公司概况', 1)
        doc.add_paragraph('芯动联科（688582.SH）是一家专注于MEMS惯性传感器芯片研发、生产和销售的高新技术企业。公司产品广泛应用于消费电子、汽车电子、工业控制等领域。')
        doc.add_paragraph(f'最新收盘价：{self.df["close"].iloc[-1]:.2f}元')
        doc.add_paragraph(f'总市值：{self.df["total_mv"].iloc[-1]/10000:.2f}亿元')
        doc.add_paragraph(f'市盈率PE(TTM)：{self.df["pe_ttm"].iloc[-1]:.2f}')
        doc.add_paragraph(f'市净率PB：{self.df["pb"].iloc[-1]:.2f}')
        doc.add_page_break()

        # 2. K线分析
        doc.add_heading('2. K线分析', 1)
        doc.add_paragraph('图1展示了芯动联科的K线走势及移动平均线分析。')
        doc.add_picture('图1_K线图与均线分析.png', width=Inches(6))
        doc.add_paragraph('数据来源：Tushare金融数据接口')

        doc.add_paragraph('\n解读：')
        current_price = self.df['close'].iloc[-1]
        ma5 = self.df['MA5'].iloc[-1]
        ma10 = self.df['MA10'].iloc[-1]
        ma20 = self.df['MA20'].iloc[-1]

        if not pd.isna(ma5):
            if current_price > ma5:
                doc.add_paragraph('• 当前股价位于MA5均线之上，短期趋势偏强。')
            else:
                doc.add_paragraph('• 当前股价位于MA5均线之下，短期趋势偏弱。')

        doc.add_page_break()

        # 3. 基本面分析
        doc.add_heading('3. 基本面分析', 1)
        doc.add_paragraph('图5展示了公司的估值指标变化。')
        doc.add_picture('图5_估值指标分析.png', width=Inches(6))

        doc.add_paragraph('\n估值分析：')
        doc.add_paragraph(f'• 当前市盈率PE为{self.df["pe"].iloc[-1]:.2f}倍，PE(TTM)为{self.df["pe_ttm"].iloc[-1]:.2f}倍')
        doc.add_paragraph(f'• 当前市净率PB为{self.df["pb"].iloc[-1]:.2f}倍')
        doc.add_paragraph(f'• 当前市销率PS为{self.df["ps"].iloc[-1]:.2f}倍')
        doc.add_paragraph(f'• 总市值为{self.df["total_mv"].iloc[-1]/10000:.2f}亿元')
        doc.add_page_break()

        # 4. 技术面分析
        doc.add_heading('4. 技术面分析', 1)

        doc.add_paragraph('4.1 MACD指标')
        doc.add_picture('图2_MACD指标分析.png', width=Inches(6))
        latest_macd = self.df['MACD_HIST'].iloc[-1]
        if latest_macd > 0:
            doc.add_paragraph(f'MACD柱状图为正数({latest_macd:.3f})，显示多头力量占优。')
        else:
            doc.add_paragraph(f'MACD柱状图为负数({latest_macd:.3f})，显示空头力量占优。')

        doc.add_paragraph('\n4.2 RSI和KDJ指标')
        doc.add_picture('图3_RSI和KDJ指标分析.png', width=Inches(6))
        latest_rsi = self.df['RSI'].iloc[-1]
        if latest_rsi > 70:
            doc.add_paragraph(f'RSI为{latest_rsi:.2f}，处于超买区间，注意回调风险。')
        elif latest_rsi < 30:
            doc.add_paragraph(f'RSI为{latest_rsi:.2f}，处于超卖区间，可能存在反弹机会。')
        else:
            doc.add_paragraph(f'RSI为{latest_rsi:.2f}，处于正常区间。')

        doc.add_paragraph('\n4.3 布林带指标')
        doc.add_picture('图4_布林带指标.png', width=Inches(6))

        doc.add_paragraph('\n4.4 换手率和量比')
        doc.add_picture('图6_换手率和量比分析.png', width=Inches(6))
        latest_turnover = self.df['turnover_rate'].iloc[-1]
        latest_volume_ratio = self.df['volume_ratio'].iloc[-1]
        doc.add_paragraph(f'最新换手率为{latest_turnover:.2f}%，量比为{latest_volume_ratio:.2f}。')

        doc.add_page_break()

        # 5. 投资建议
        doc.add_heading('5. 投资建议', 1)
        doc.add_paragraph('综合以上分析，给出以下投资建议：')
        doc.add_paragraph('• 建议投资者密切关注技术指标变化')
        doc.add_paragraph('• 注意估值水平是否合理')
        doc.add_paragraph('• 关注行业政策和市场情绪变化')
        doc.add_paragraph('• 控制仓位，做好风险管理')
        doc.add_page_break()

        # 6. 风险提示
        doc.add_heading('6. 风险提示', 1)
        doc.add_paragraph('• 股市有风险，投资需谨慎')
        doc.add_paragraph('• 本报告仅供参考，不构成投资建议')
        doc.add_paragraph('• 技术指标存在滞后性，需结合其他因素综合判断')
        doc.add_paragraph('• 市场环境和政策变化可能影响股价走势')

        # 保存文档
        doc.save('芯动联科股票分析报告_完整版.docx')
        print("✓ Word报告生成完成：芯动联科股票分析报告_完整版.docx\n")

    def generate_jupyter_notebook(self):
        """生成Jupyter Notebook"""
        print("开始生成Jupyter Notebook...")

        notebook_content = {
            "cells": [
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": [
                        "# 芯动联科（688582.SH）股票分析\n",
                        "\n",
                        "本Notebook包含完整的数据分析和可视化代码。"
                    ]
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": [
                        "import json\n",
                        "import pandas as pd\n",
                        "import numpy as np\n",
                        "import matplotlib.pyplot as plt\n",
                        "import warnings\n",
                        "warnings.filterwarnings('ignore')\n",
                        "\n",
                        "# 设置中文字体\n",
                        "plt.rcParams['font.sans-serif'] = ['SimSun', 'DejaVu Sans']\n",
                        "plt.rcParams['axes.unicode_minus'] = False\n",
                        "\n",
                        "print('库导入成功！')"
                    ]
                },
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": [
                        "## 1. 读取数据"
                    ]
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": [
                        "# 读取股票数据\n",
                        "with open('stock_data.json', 'r', encoding='utf-8') as f:\n",
                        "    stock_data = json.load(f)\n",
                        "\n",
                        "# 读取基本面数据\n",
                        "with open('daily_basic_data.json', 'r', encoding='utf-8') as f:\n",
                        "    basic_data = json.load(f)\n",
                        "\n",
                        "# 转换为DataFrame\n",
                        "df_stock = pd.DataFrame(stock_data)\n",
                        "df_basic = pd.DataFrame(basic_data)\n",
                        "\n",
                        "print(f'股票数据条数: {len(df_stock)}')\n",
                        "print(f'基本面数据条数: {len(df_basic)}')"
                    ]
                },
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": [
                        "## 2. 技术指标计算"
                    ]
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": [
                        "# 数据处理\n",
                        "df_stock['trade_date'] = pd.to_datetime(df_stock['trade_date'], format='%Y%m%d')\n",
                        "df_basic['trade_date'] = pd.to_datetime(df_basic['trade_date'], format='%Y%m%d')\n",
                        "\n",
                        "df_stock = df_stock.sort_values('trade_date').reset_index(drop=True)\n",
                        "df_basic = df_basic.sort_values('trade_date').reset_index(drop=True)\n",
                        "\n",
                        "# 合并数据\n",
                        "df = pd.merge(df_stock, df_basic, on='trade_date', how='left')\n",
                        "\n",
                        "# 计算移动平均线\n",
                        "df['MA5'] = df['close'].rolling(window=5).mean()\n",
                        "df['MA10'] = df['close'].rolling(window=10).mean()\n",
                        "df['MA20'] = df['close'].rolling(window=20).mean()\n",
                        "df['MA60'] = df['close'].rolling(window=60).mean()\n",
                        "\n",
                        "# 计算MACD\n",
                        "exp1 = df['close'].ewm(span=12, adjust=False).mean()\n",
                        "exp2 = df['close'].ewm(span=26, adjust=False).mean()\n",
                        "df['MACD_DIF'] = exp1 - exp2\n",
                        "df['MACD_DEA'] = df['MACD_DIF'].ewm(span=9, adjust=False).mean()\n",
                        "df['MACD_HIST'] = (df['MACD_DIF'] - df['MACD_DEA']) * 2\n",
                        "\n",
                        "print('技术指标计算完成！')\n",
                        "print(df[['trade_date', 'close', 'MA5', 'MA10', 'MA20', 'MACD_DIF']].tail())"
                    ]
                },
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": [
                        "## 3. 可视化分析"
                    ]
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": [
                        "# 绘制K线图和均线\n",
                        "import matplotlib.dates as mdates\n",
                        "\n",
                        "fig, ax = plt.subplots(figsize=(14, 8))\n",
                        "\n",
                        "# 绘制K线\n",
                        "for i in range(len(df)):\n",
                        "    date = df.iloc[i]['trade_date']\n",
                        "    open_price = df.iloc[i]['open']\n",
                        "    close_price = df.iloc[i]['close']\n",
                        "    high_price = df.iloc[i]['high']\n",
                        "    low_price = df.iloc[i]['low']\n",
                        "    \n",
                        "    color = 'red' if close_price >= open_price else 'green'\n",
                        "    ax.plot([date, date], [low_price, high_price], color=color, linewidth=1, alpha=0.8)\n",
                        "    ax.plot([date, date], [open_price, close_price], color=color, linewidth=4, alpha=0.8)\n",
                        "\n",
                        "# 绘制均线\n",
                        "ax.plot(df['trade_date'], df['MA5'], label='MA5', linewidth=1.5)\n",
                        "ax.plot(df['trade_date'], df['MA10'], label='MA10', linewidth=1.5)\n",
                        "ax.plot(df['trade_date'], df['MA20'], label='MA20', linewidth=1.5)\n",
                        "ax.plot(df['trade_date'], df['MA60'], label='MA60', linewidth=1.5)\n",
                        "\n",
                        "ax.set_title('芯动联科（688582.SH）- K线图与均线分析', fontsize=16, fontweight='bold')\n",
                        "ax.set_xlabel('日期')\n",
                        "ax.set_ylabel('价格（元）')\n",
                        "ax.legend()\n",
                        "ax.grid(True, alpha=0.3)\n",
                        "plt.xticks(rotation=45)\n",
                        "plt.tight_layout()\n",
                        "plt.show()"
                    ]
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": [
                        "# 绘制MACD指标\n",
                        "fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10), height_ratios=[2, 1])\n",
                        "\n",
                        "ax1.plot(df['trade_date'], df['close'], color='black', linewidth=2, label='收盘价')\n",
                        "ax1.set_title('MACD指标分析', fontsize=16, fontweight='bold')\n",
                        "ax1.legend()\n",
                        "ax1.grid(True, alpha=0.3)\n",
                        "\n",
                        "ax2.plot(df['trade_date'], df['MACD_DIF'], label='DIF', linewidth=1.5)\n",
                        "ax2.plot(df['trade_date'], df['MACD_DEA'], label='DEA', linewidth=1.5)\n",
                        "colors = ['red' if x > 0 else 'green' for x in df['MACD_HIST']]\n",
                        "ax2.bar(df['trade_date'], df['MACD_HIST'], color=colors, alpha=0.5, label='MACD')\n",
                        "ax2.set_xlabel('日期')\n",
                        "ax2.legend()\n",
                        "ax2.grid(True, alpha=0.3)\n",
                        "plt.xticks(rotation=45)\n",
                        "plt.tight_layout()\n",
                        "plt.show()"
                    ]
                },
                {
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": [
                        "## 4. 数据分析结论"
                    ]
                },
                {
                    "cell_type": "code",
                    "execution_count": None,
                    "metadata": {},
                    "outputs": [],
                    "source": [
                        "# 输出关键指标\n",
                        "print('=== 芯动联科（688582.SH）关键指标 ===')\n",
                        "print(f'最新收盘价: {df[\"close\"].iloc[-1]:.2f}元')\n",
                        "print(f'最新PE: {df[\"pe\"].iloc[-1]:.2f}')\n",
                        "print(f'最新PB: {df[\"pb\"].iloc[-1]:.2f}')\n",
                        "print(f'最新PS: {df[\"ps\"].iloc[-1]:.2f}')\n",
                        "print(f'总市值: {df[\"total_mv\"].iloc[-1]/10000:.2f}亿元')\n",
                        "print(f'换手率: {df[\"turnover_rate\"].iloc[-1]:.2f}%')\n",
                        "print(f'量比: {df[\"volume_ratio\"].iloc[-1]:.2f}')"
                    ]
                }
            ],
            "metadata": {
                "kernelspec": {
                    "display_name": "Python 3",
                    "language": "python",
                    "name": "python3"
                },
                "language_info": {
                    "name": "python",
                    "version": "3.13.12"
                }
            },
            "nbformat": 4,
            "nbformat_minor": 5
        }

        with open('芯动联科股票分析_完整版.ipynb', 'w', encoding='utf-8') as f:
            json.dump(notebook_content, f, ensure_ascii=False, indent=2)

        print("✓ Jupyter Notebook生成完成：芯动联科股票分析_完整版.ipynb\n")

    def run_complete_analysis(self):
        """运行完整分析流程"""
        print("="*60)
        print("芯动联科（688582.SH）股票完整分析")
        print("="*60)

        # 生成技术指标
        self.generate_all_indicators()

        # 生成图表
        self.generate_all_charts()

        # 生成Word报告
        try:
            self.generate_word_report()
        except Exception as e:
            print(f"Word报告生成失败：{e}")
            print("跳过Word报告生成...")

        # 生成Jupyter Notebook
        self.generate_jupyter_notebook()

        print("="*60)
        print("✓ 所有任务完成！")
        print("="*60)
        print("\n生成的文件：")
        print("1. 图1_K线图与均线分析.png")
        print("2. 图2_MACD指标分析.png")
        print("3. 图3_RSI和KDJ指标分析.png")
        print("4. 图4_布林带指标.png")
        print("5. 图5_估值指标分析.png")
        print("6. 图6_换手率和量比分析.png")
        print("7. 芯动联科股票分析报告_完整版.docx")
        print("8. 芯动联科股票分析_完整版.ipynb")
        print("\n所有文件已保存到工作目录。")

if __name__ == '__main__':
    # 创建分析器实例
    analyzer = StockAnalyzer(
        stock_data_file='stock_data.json',
        basic_data_file='daily_basic_data.json'
    )

    # 运行完整分析
    analyzer.run_complete_analysis()
